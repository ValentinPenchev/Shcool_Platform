import re
import io
from docx import Document

# Символи, които не се набират от стандартна клавиатура (кирилица/латиница/цифри/
# основна пунктуация) - използват се за откриване на "специални символи" (Insert > Symbol)
_ALLOWED_CHARS_PATTERN = re.compile(
    r'^[a-zA-Zа-яА-ЯёЁ0-9\s\.\,\!\?\-\:\;\'\"\(\)\/]*$'
)


def _iter_runs(doc):
    for p in doc.paragraphs:
        for r in p.runs:
            yield r
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        yield r


def _has_paragraph_alignment(doc):
    return any(p.alignment is not None for p in doc.paragraphs)


def _has_font_name(doc):
    return any(r.font.name for r in _iter_runs(doc))


def _has_font_size(doc):
    return any(r.font.size is not None for r in _iter_runs(doc))


def _has_font_color(doc):
    for r in _iter_runs(doc):
        try:
            if r.font.color and r.font.color.rgb is not None:
                return True
        except Exception:
            continue
    return False


def _has_bold(doc):
    return any(r.font.bold for r in _iter_runs(doc))


def _has_italic(doc):
    return any(r.font.italic for r in _iter_runs(doc))


def _has_underline(doc):
    return any(r.font.underline for r in _iter_runs(doc))


def _has_special_characters(doc):
    for p in doc.paragraphs:
        text = p.text
        if text and not _ALLOWED_CHARS_PATTERN.match(text):
            return True
    return False


def _has_merged_cells(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc_xml = cell._tc.xml
                if "gridSpan" in tc_xml or "vMerge" in tc_xml:
                    return True
    return False


def _has_cell_shading(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc_xml = cell._tc.xml
                if "<w:shd" not in tc_xml:
                    continue
                match = re.search(r'w:fill="([0-9A-Fa-f]{6}|auto)"', tc_xml)
                if match and match.group(1).lower() not in ("auto", "ffffff"):
                    return True
    return False


def _has_image(doc):
    return any("image" in rel.target_ref for rel in doc.part.rels.values())


def evaluate_word(file_bytes: bytes, criteria: dict) -> dict:
    """Оценява .docx файл спрямо предоставените критерии."""
    doc = Document(io.BytesIO(file_bytes))
    score = 0
    max_score = 0
    details = []

    def add_check(key, label, passed_note, failed_note, passed):
        nonlocal score, max_score
        rule = criteria[key]
        rule_points = rule.get("points", 1) if isinstance(rule, dict) else 1
        max_score += rule_points
        if passed:
            score += rule_points
        details.append({
            "criterion": label,
            "passed": passed,
            "score": rule_points if passed else 0,
            "max": rule_points,
            "note": passed_note if passed else failed_note
        })

    if "font" in criteria:
        has_text = any(p.text.strip() for p in doc.paragraphs)
        add_check("font", "Форматиране / Текст",
                   "Намерен е съдържателен текст в документа.",
                   "Документът е празен или липсва текст.", has_text)

    if "table" in criteria:
        add_check("table", "Наличие на таблица",
                   f"Намерена е таблица в документа (общо: {len(doc.tables)}).",
                   "Не е намерена таблица в документа.", len(doc.tables) > 0)

    if "image" in criteria:
        add_check("image", "Вмъкнато изображение",
                   "Намерено е изображение в документа.",
                   "Липсва вмъкнато изображение.", _has_image(doc))

    if "paragraph_alignment" in criteria:
        add_check("paragraph_alignment", "Подравняване на абзац",
                   "Намерено е зададено подравняване на поне един абзац.",
                   "Не е зададено подравняване на абзаците.", _has_paragraph_alignment(doc))

    if "font_name" in criteria:
        add_check("font_name", "Шрифт",
                   "Зададен е конкретен шрифт на текста.",
                   "Не е зададен конкретен шрифт.", _has_font_name(doc))

    if "font_size" in criteria:
        add_check("font_size", "Големина на шрифта",
                   "Зададена е конкретна големина на шрифта.",
                   "Не е зададена конкретна големина на шрифта.", _has_font_size(doc))

    if "font_color" in criteria:
        add_check("font_color", "Цвят на текста",
                   "Зададен е цвят на текста.",
                   "Не е зададен цвят на текста.", _has_font_color(doc))

    if "bold" in criteria:
        add_check("bold", "Удебеляване",
                   "Намерен е удебелен текст.",
                   "Не е намерен удебелен текст.", _has_bold(doc))

    if "italic" in criteria:
        add_check("italic", "Курсив",
                   "Намерен е текст в курсив.",
                   "Не е намерен текст в курсив.", _has_italic(doc))

    if "underline" in criteria:
        add_check("underline", "Подчертаване",
                   "Намерен е подчертан текст.",
                   "Не е намерен подчертан текст.", _has_underline(doc))

    if "special_characters" in criteria:
        add_check("special_characters", "Специални символи",
                   "Намерени са специални символи в текста.",
                   "Не са намерени специални символи.", _has_special_characters(doc))

    if "merged_cells" in criteria:
        add_check("merged_cells", "Обединени клетки в таблица",
                   "Намерени са обединени клетки в таблица.",
                   "Не са намерени обединени клетки в таблица.", _has_merged_cells(doc))

    if "cell_shading" in criteria:
        add_check("cell_shading", "Оцветяване на клетка",
                   "Намерена е оцветена клетка в таблица.",
                   "Не е намерена оцветена клетка в таблица.", _has_cell_shading(doc))

    percentage = round((score / max_score) * 100) if max_score > 0 else 0
    return {
        "score": score,
        "max_score": max_score,
        "percentage": percentage,
        "details": details
    }
